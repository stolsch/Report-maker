# === app.py (Streamlit UI with OpenAI SDK v1+) ===

import streamlit as st
import requests
from bs4 import BeautifulSoup
from openai import OpenAI
from docx import Document
from datetime import datetime
import os

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

TEMPLATE_PATH = "CyberSec Target Profile_Template_1744109169569.docx"


def search_web(query):
    url = f"https://www.google.com/search?q={query}"
    headers = {"User-Agent": "Mozilla/5.0"}
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.text, 'html.parser')
    links = [a['href'] for a in soup.select("a") if a['href'].startswith("http")]
    return links[:5]


def summarize_sections(company_name):
    prompt = f"""
    Fill the following cybersecurity company profile template with real or best-estimate data about a company called: {company_name}.
    Include: Company Overview, Team, Business Model, Technology, Market Position, Cybersecurity Services, Financial Info.
    Also explain common terms like MVP, SaaS, AI, etc., at the end as a glossary.
    Format responses in structured paragraphs with each section clearly labeled.
    """
    chat_completion = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7
    )
    return chat_completion.choices[0].message.content


def generate_profile(company_name):
    doc = Document(TEMPLATE_PATH)
    summary = summarize_sections(company_name)
    sources = search_web(company_name + " cybersecurity")

    doc.add_page_break()
    doc.add_heading(f"Auto-Generated Profile: {company_name}", level=1)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(summary)
    doc.add_heading("Sources", level=1)
    for link in sources:
        doc.add_paragraph(link)

    output_path = f"{company_name.replace(' ', '_')}_CyberProfile.docx"
    doc.save(output_path)
    return output_path


# === Streamlit UI ===
st.set_page_config(page_title="Cyber Startup Reporter")
st.title("Cybersecurity Startup Profile Generator")

company_name = st.text_input("Enter Company Name")

if st.button("Generate Report") and company_name:
    with st.spinner("Generating profile, please wait..."):
        filepath = generate_profile(company_name)
        with open(filepath, "rb") as file:
            st.success("Profile generated!")
            st.download_button(label="Download Report", data=file, file_name=filepath, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
